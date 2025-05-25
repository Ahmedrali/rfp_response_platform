"""
Background tasks for document processing with RAG integration
"""
import os
import traceback
from datetime import datetime
from typing import List, Dict, Any, Optional
from uuid import UUID
import PyPDF2
import docx
import structlog
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.celery_app import celery_app
from app.models.company_document import CompanyDocument, ProcessingStatus
from app.models.document_chunk import DocumentChunk
from app.services.rag_document_processor import DocumentProcessor
from app.utils.database import get_async_session_factory

log = structlog.get_logger(__name__)


async def get_db_session():
    """Get async database session"""
    async_session = get_async_session_factory()
    async with async_session() as session:
        yield session


@celery_app.task(name="process_document", bind=True, max_retries=3)
def process_document(self, document_id: str, company_id: str):
    """
    Process a document using RAG document processor
    
    Args:
        document_id: Document ID (UUID string)
        company_id: Company ID (UUID string)
    """
    log.info("Starting document processing",
             document_id=document_id,
             company_id=company_id,
             task_id=self.request.id)
    
    try:
        # Run the async processing within a sync context
        import asyncio
        loop = asyncio.get_event_loop()
        return loop.run_until_complete(_process_document_async(document_id, company_id))
    
    except Exception as e:
        log.error("Document processing failed",
                 document_id=document_id,
                 company_id=company_id,
                 error=str(e),
                 traceback=traceback.format_exc())
        
        # Retry the task if appropriate
        raise self.retry(exc=e, countdown=60)  # Retry after 1 minute


async def _process_document_async(document_id: str, company_id: str) -> Dict[str, Any]:
    """
    Async implementation of document processing
    
    Args:
        document_id: Document ID (UUID string)
        company_id: Company ID (UUID string)
        
    Returns:
        Processing result details
    """
    # Initialize RAG document processor
    processor = DocumentProcessor()
    
    # Create database session
    async_session = get_async_session_factory()
    async with async_session() as db:
        
        try:
            # 1. Get document from database
            doc_query = select(CompanyDocument).where(
                CompanyDocument.id == UUID(document_id),
                CompanyDocument.company_id == UUID(company_id)
            )
            result = await db.execute(doc_query)
            document = result.scalar_one_or_none()
            
            if not document:
                log.error("Document not found",
                         document_id=document_id,
                         company_id=company_id)
                return {"status": "error", "message": "Document not found"}
            
            # 2. Update document status to PROCESSING
            document.processing_status = ProcessingStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            await db.commit()
            
            # 3. Extract text from document
            file_path = document.file_path
            log.info("Extracting text from document",
                    document_id=document_id,
                    file_path=file_path)
            
            extracted_text = await extract_text_from_file(file_path)
            if not extracted_text:
                document.processing_status = ProcessingStatus.FAILED
                document.error_message = "Failed to extract text from document"
                document.processing_completed_at = datetime.utcnow()
                await db.commit()
                log.error("Text extraction failed",
                         document_id=document_id,
                         file_path=file_path)
                return {"status": "error", "message": "Text extraction failed"}
            
            # 4. Process document with RAG processor
            log.info("Processing document with RAG",
                    document_id=document_id,
                    company_id=company_id)
            
            # Use company_id as db_identifier for company-specific isolation
            chunk_ids = processor.index_document(
                file_content=extracted_text,
                db_identifier=str(company_id)
            )
            
            # 5. Store chunk IDs in database
            log.info("Storing chunk information",
                    document_id=document_id,
                    chunk_count=len(chunk_ids))
            
            # Create chunk records
            chunks_to_add = []
            for idx, chunk_id in enumerate(chunk_ids):
                chunk = DocumentChunk(
                    document_id=UUID(document_id),
                    chunk_index=idx,
                    chunk_id=chunk_id,
                    meta_data={
                        "processed_at": datetime.utcnow().isoformat()
                    }
                )
                chunks_to_add.append(chunk)
            
            # Bulk insert all chunks at once for better performance
            if chunks_to_add:
                db.add_all(chunks_to_add)
            
            # Update document with chunk count and status
            document.chunk_count = len(chunk_ids)
            document.processing_status = ProcessingStatus.COMPLETED
            document.processing_completed_at = datetime.utcnow()
            meta_data_dict = {}
            if document.meta_data:
                meta_data_dict = document.meta_data
            
            meta_data_dict["processing_details"] = {
                "chunk_count": len(chunk_ids),
                "processed_at": datetime.utcnow().isoformat()
            }
            
            document.meta_data = meta_data_dict
            
            await db.commit()
            
            log.info("Document processing completed successfully",
                    document_id=document_id,
                    company_id=company_id,
                    chunk_count=len(chunk_ids))
            
            return {
                "status": "success",
                "document_id": document_id,
                "company_id": company_id,
                "chunk_count": len(chunk_ids),
                "chunk_ids": chunk_ids
            }
            
        except Exception as e:
            await db.rollback()
            
            # Update document status to FAILED
            try:
                doc_query = select(CompanyDocument).where(
                    CompanyDocument.id == UUID(document_id),
                    CompanyDocument.company_id == UUID(company_id)
                )
                result = await db.execute(doc_query)
                document = result.scalar_one_or_none()
                
                if document:
                    document.processing_status = ProcessingStatus.FAILED
                    document.error_message = str(e)
                    document.processing_completed_at = datetime.utcnow()
                    await db.commit()
            except Exception as inner_e:
                log.error("Failed to update document status after error",
                         document_id=document_id,
                         error=str(inner_e))
            
            log.error("Document processing failed with exception",
                     document_id=document_id,
                     company_id=company_id,
                     error=str(e),
                     traceback=traceback.format_exc())
            
            # Re-raise for retry
            raise


async def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text content from document files
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        if not os.path.exists(file_path):
            log.error(f"File not found: {file_path}")
            return None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # PDF extraction
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        
        # Word document extraction
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        
        # Plain text
        elif file_extension in ['.txt', '.md', '.csv']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        else:
            log.error(f"Unsupported file type: {file_extension}")
            return None
            
    except Exception as e:
        log.error(f"Text extraction failed: {str(e)}")
        return None


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files"""
    text = ""
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
                
        return text.strip()
    except Exception as e:
        log.error(f"PDF extraction failed: {str(e)}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Include text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return '\n'.join(full_text)
    except Exception as e:
        log.error(f"DOCX extraction failed: {str(e)}")
        raise
